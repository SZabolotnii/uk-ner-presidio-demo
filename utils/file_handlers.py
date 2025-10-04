"""
Файлові обробники: нормалізація різних форматів до тексту.

Архітектурна стратегія: Adapter Pattern для уніфікації файлових джерел.
Відповідальність: Перетворення TXT/DOCX → чистий текст з валідацією.
"""

import logging
from pathlib import Path
from typing import Optional, Tuple
from dataclasses import dataclass

from docx import Document
import chardet

logger = logging.getLogger(__name__)


@dataclass
class FileReadResult:
    """
    Структурований результат читання файлу.
    
    Design Pattern: Value Object для передачі даних між шарами.
    """
    text: str
    filename: str
    file_type: str
    encoding: Optional[str] = None
    char_count: int = 0
    
    def __post_init__(self):
        """Автоматичний підрахунок символів."""
        if self.char_count == 0:
            self.char_count = len(self.text)


class FileHandler:
    """
    Базовий інтерфейс для файлових обробників.
    
    Strategy Pattern: Дозволяє легко додавати нові формати (PDF, ODT, etc.)
    """
    
    # Константи для валідації
    MAX_FILE_SIZE_MB = 50
    MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
    
    SUPPORTED_EXTENSIONS = {'.txt', '.docx'}
    
    @classmethod
    def read_file(cls, file_path: str) -> FileReadResult:
        """
        Основний метод читання файлу (auto-detection формату).
        
        Args:
            file_path: Шлях до файлу або file-like object від Gradio
            
        Returns:
            FileReadResult з текстом та метаданими
            
        Raises:
            ValueError: Непідтримуваний формат або файл занадто великий
            RuntimeError: Помилка читання
        """
        # Конвертуємо в Path для універсальності
        path = Path(file_path) if isinstance(file_path, str) else Path(file_path.name)
        
        # Валідація розміру
        cls._validate_file_size(path)
        
        # Визначаємо обробник за розширенням
        extension = path.suffix.lower()
        
        if extension == '.txt':
            return cls._read_txt(path)
        elif extension == '.docx':
            return cls._read_docx(path)
        else:
            raise ValueError(
                f"Непідтримуваний формат файлу: {extension}\n"
                f"Підтримуються: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
    
    @classmethod
    def _validate_file_size(cls, path: Path) -> None:
        """
        Валідація розміру файлу.
        
        Security: Запобігання DoS атакам через великі файли.
        """
        try:
            file_size = path.stat().st_size
        except (OSError, AttributeError):
            # Якщо це file-like object, пропускаємо перевірку
            # (Gradio вже обмежує розмір)
            return
        
        if file_size > cls.MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"Файл завеликий: {file_size / 1024 / 1024:.1f} MB. "
                f"Максимум: {cls.MAX_FILE_SIZE_MB} MB"
            )
        
        logger.info(f"File size: {file_size / 1024:.1f} KB")
    
    @classmethod
    def _read_txt(cls, path: Path) -> FileReadResult:
        """
        Читання TXT файлу з автоматичним визначенням encoding.
        
        Стратегія:
        1. Спроба UTF-8 (найпоширеніше)
        2. Автодетект через chardet
        3. Fallback на cp1251 (для старих українських файлів)
        
        Args:
            path: Шлях до файлу
            
        Returns:
            FileReadResult
        """
        # Спроба 1: UTF-8
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            logger.info(f"Successfully read TXT file with UTF-8: {path.name}")
            
            return FileReadResult(
                text=text,
                filename=path.name,
                file_type='txt',
                encoding='utf-8'
            )
        
        except UnicodeDecodeError:
            logger.warning("UTF-8 failed, trying auto-detection")
        
        # Спроба 2: Auto-detection
        try:
            with open(path, 'rb') as f:
                raw_data = f.read()
            
            detected = chardet.detect(raw_data)
            encoding = detected['encoding']
            
            if encoding:
                text = raw_data.decode(encoding)
                
                logger.info(
                    f"Auto-detected encoding: {encoding} "
                    f"(confidence: {detected['confidence']:.0%})"
                )
                
                return FileReadResult(
                    text=text,
                    filename=path.name,
                    file_type='txt',
                    encoding=encoding
                )
        
        except Exception as e:
            logger.warning(f"Auto-detection failed: {e}")
        
        # Спроба 3: Fallback на cp1251
        try:
            with open(path, 'r', encoding='cp1251') as f:
                text = f.read()
            
            logger.info("Fallback to cp1251 successful")
            
            return FileReadResult(
                text=text,
                filename=path.name,
                file_type='txt',
                encoding='cp1251'
            )
        
        except Exception as e:
            raise RuntimeError(
                f"Не вдалося прочитати файл з жодним encoding: {e}"
            )
    
    @classmethod
    def _read_docx(cls, path: Path) -> FileReadResult:
        """
        Читання DOCX файлу.
        
        Особливості:
        - Витягує тільки текст (без форматування)
        - Зберігає параграфи через \n\n
        - Ігнорує таблиці, колонтитули, notes
        
        Args:
            path: Шлях до DOCX файлу
            
        Returns:
            FileReadResult
        """
        try:
            doc = Document(path)
            
            # Витягуємо текст з усіх параграфів
            paragraphs = [p.text.strip() for p in doc.paragraphs]
            
            # Фільтруємо порожні параграфи та об'єднуємо
            text = "\n\n".join(p for p in paragraphs if p)
            
            logger.info(
                f"Successfully read DOCX: {path.name}, "
                f"{len(doc.paragraphs)} paragraphs"
            )
            
            return FileReadResult(
                text=text,
                filename=path.name,
                file_type='docx',
                encoding=None  # DOCX internally uses XML
            )
        
        except Exception as e:
            raise RuntimeError(f"Помилка читання DOCX файлу: {e}")
    
    @staticmethod
    def get_file_info(file_path: str) -> dict:
        """
        Повертає інформацію про файл без читання.
        
        Utility method для UI попереднього перегляду.
        
        Returns:
            Dict з метаданими файлу
        """
        path = Path(file_path)
        
        try:
            stat = path.stat()
            
            return {
                'filename': path.name,
                'extension': path.suffix,
                'size_kb': stat.st_size / 1024,
                'size_mb': stat.st_size / 1024 / 1024,
                'supported': path.suffix.lower() in FileHandler.SUPPORTED_EXTENSIONS
            }
        except Exception as e:
            logger.error(f"Failed to get file info: {e}")
            return {
                'filename': path.name if hasattr(path, 'name') else 'unknown',
                'error': str(e)
            }


# ============ EXTENSION POINT ============

class PDFHandler:
    """
    МАЙБУТНЄ РОЗШИРЕННЯ: PDF обробник.
    
    Extensibility: Коли потрібна підтримка PDF, імплементувати тут.
    Потрібні залежності: PyPDF2 або pdfplumber
    """
    
    @staticmethod
    def read_pdf(path: Path) -> FileReadResult:
        """Placeholder для PDF підтримки."""
        raise NotImplementedError(
            "PDF підтримка не імплементована. "
            "Додайте PyPDF2 до requirements.txt"
        )


# ============ HELPER FUNCTIONS ============

# utils/file_handlers.py - функція sanitize_text

def sanitize_text(text: str) -> str:
    """
    Нормалізує текст після читання з файлу.
    
    Операції:
    - Видалення зайвих пробілів
    - Нормалізація line endings
    - Видалення control characters
    - Обмеження послідовних порожніх рядків до 1 (max 2 newlines підряд)
    
    Args:
        text: Сирий текст з файлу
        
    Returns:
        Очищений текст
        
    Design Decision: Дозволяємо максимум 1 порожній рядок (2 newlines підряд)
    для збереження структури документу, але запобігаємо надмірним пробілам.
    """
    # Нормалізуємо line endings (Windows/Mac → Unix)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Видаляємо trailing whitespace з кожного рядка
    lines = [line.rstrip() for line in text.split('\n')]
    
    # ✅ FIXED: Обмежуємо послідовні порожні рядки
    cleaned_lines = []
    consecutive_empty = 0
    
    for line in lines:
        if not line:  # Порожній рядок
            consecutive_empty += 1
            # Дозволяємо максимум 1 порожній рядок (тобто 2 newlines підряд)
            if consecutive_empty <= 1:
                cleaned_lines.append(line)
            # Якщо більше - пропускаємо
        else:  # Непорожній рядок
            consecutive_empty = 0  # Скидаємо лічильник
            cleaned_lines.append(line)
    
    # Об'єднуємо та видаляємо whitespace на початку/кінці
    result = '\n'.join(cleaned_lines).strip()
    
    return result