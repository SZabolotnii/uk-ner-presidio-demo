"""
Експорт результатів деідентифікації в різні формати.

Архітектурна стратегія: Strategy Pattern для підтримки множини форматів.
Extensibility: Легке додавання нових форматів експорту (PDF, HTML, etc.)
"""

import logging
import json
from pathlib import Path
from datetime import datetime
from typing import List, Optional
from dataclasses import asdict

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.analyzer import AnalysisResult
from presidio_analyzer import RecognizerResult

logger = logging.getLogger(__name__)


class ExportFormat:
    """Константи підтримуваних форматів експорту."""
    TXT = 'txt'
    DOCX = 'docx'
    JSON = 'json'
    CSV = 'csv'
    MARKDOWN = 'md'


class FileExporter:
    """
    Універсальний експортер результатів аналізу.
    
    Design Pattern: Facade + Strategy для різних форматів.
    Responsibility: Конвертація AnalysisResult → файлові формати.
    """
    
    @staticmethod
    def export_anonymized_text(
        result: AnalysisResult,
        format: str = ExportFormat.TXT,
        include_metadata: bool = True
    ) -> bytes:
        """
        Експортує анонімізований текст у вказаному форматі.
        
        Args:
            result: Результат аналізу
            format: Формат експорту (txt/docx/md)
            include_metadata: Чи додавати метадані на початок
            
        Returns:
            Байти файлу для завантаження
        """
        if format == ExportFormat.TXT:
            return FileExporter._export_txt(result, include_metadata)
        elif format == ExportFormat.DOCX:
            return FileExporter._export_docx(result, include_metadata)
        elif format == ExportFormat.MARKDOWN:
            return FileExporter._export_markdown(result, include_metadata)
        else:
            raise ValueError(f"Непідтримуваний формат: {format}")
    
    @staticmethod
    def export_entities_report(
        result: AnalysisResult,
        format: str = ExportFormat.JSON
    ) -> bytes:
        """
        Експортує звіт про знайдені сутності.
        
        Args:
            result: Результат аналізу
            format: Формат експорту (json/csv/txt)
            
        Returns:
            Байти файлу для завантаження
        """
        if format == ExportFormat.JSON:
            return FileExporter._export_entities_json(result)
        elif format == ExportFormat.CSV:
            return FileExporter._export_entities_csv(result)
        elif format == ExportFormat.TXT:
            return FileExporter._export_entities_txt(result)
        else:
            raise ValueError(f"Непідтримуваний формат: {format}")
    
    @staticmethod
    def export_full_report(
        result: AnalysisResult,
        format: str = ExportFormat.DOCX
    ) -> bytes:
        """
        Експортує повний звіт (текст + сутності + статистика).
        
        Args:
            result: Результат аналізу
            format: Формат (docx/md/txt)
            
        Returns:
            Байти файлу для завантаження
        """
        if format == ExportFormat.DOCX:
            return FileExporter._export_full_report_docx(result)
        elif format == ExportFormat.MARKDOWN:
            return FileExporter._export_full_report_md(result)
        else:
            return FileExporter._export_full_report_txt(result)
    
    # ============ TXT EXPORTERS ============
    
    @staticmethod
    def _export_txt(result: AnalysisResult, include_metadata: bool) -> bytes:
        """Експорт анонімізованого тексту в TXT."""
        content_parts = []
        
        if include_metadata:
            metadata = FileExporter._generate_metadata_header(result)
            content_parts.append(metadata)
            content_parts.append("=" * 60)
            content_parts.append("")
        
        content_parts.append(result.anonymized_text)
        
        content = "\n".join(content_parts)
        return content.encode('utf-8')
    
    @staticmethod
    def _export_entities_txt(result: AnalysisResult) -> bytes:
        """Експорт списку сутностей в TXT."""
        lines = [
            "ЗВІТ ПРО ВИЯВЛЕНІ ПЕРСОНАЛЬНІ ДАНІ",
            "=" * 60,
            "",
            f"Загальна кількість сутностей: {result.entities_count}",
            f"Дата аналізу: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "=" * 60,
            ""
        ]
        
        if result.entities_count == 0:
            lines.append("✓ Персональних даних не виявлено")
        else:
            # Групуємо за типами
            entities_by_type = {}
            for entity in result.entities:
                if entity.entity_type not in entities_by_type:
                    entities_by_type[entity.entity_type] = []
                entities_by_type[entity.entity_type].append(entity)
            
            for entity_type, entities in sorted(entities_by_type.items()):
                lines.append(f"\n📌 {entity_type} ({len(entities)} знайдено)")
                lines.append("-" * 40)
                
                for idx, entity in enumerate(sorted(entities, key=lambda x: x.start), 1):
                    entity_text = result.original_text[entity.start:entity.end]
                    lines.append(
                        f"{idx}. '{entity_text}' "
                        f"[позиція {entity.start}-{entity.end}, "
                        f"впевненість {entity.score:.0%}]"
                    )
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    @staticmethod
    def _export_full_report_txt(result: AnalysisResult) -> bytes:
        """Повний звіт у TXT."""
        lines = [
            "ПОВНИЙ ЗВІТ ДЕІДЕНТИФІКАЦІЇ",
            "=" * 60,
            "",
            FileExporter._generate_metadata_header(result),
            "",
            "=" * 60,
            "",
            "АНОНІМІЗОВАНИЙ ТЕКСТ:",
            "-" * 60,
            result.anonymized_text,
            "",
            "=" * 60,
            ""
        ]
        
        # Додаємо звіт про сутності
        entities_report = FileExporter._export_entities_txt(result).decode('utf-8')
        lines.append(entities_report)
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    # ============ JSON EXPORTERS ============
    
    @staticmethod
    def _export_entities_json(result: AnalysisResult) -> bytes:
        """Експорт сутностей у JSON (машинно-читабельний формат)."""
        data = {
            "metadata": {
                "analysis_timestamp": datetime.now().isoformat(),
                "total_entities": result.entities_count,
                "original_text_length": len(result.original_text)
            },
            "entities": [
                {
                    "type": entity.entity_type,
                    "text": result.original_text[entity.start:entity.end],
                    "start": entity.start,
                    "end": entity.end,
                    "confidence": round(entity.score, 3)
                }
                for entity in sorted(result.entities, key=lambda x: x.start)
            ],
            "statistics": FileExporter._calculate_statistics(result)
        }
        
        json_str = json.dumps(data, ensure_ascii=False, indent=2)
        return json_str.encode('utf-8')
    
    # ============ CSV EXPORTERS ============
    
    @staticmethod
    def _export_entities_csv(result: AnalysisResult) -> bytes:
        """Експорт сутностей у CSV (для аналізу в Excel)."""
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.writer(output)
        
        # Header
        writer.writerow([
            "Тип сутності", 
            "Текст", 
            "Початок", 
            "Кінець", 
            "Впевненість (%)"
        ])
        
        # Дані
        for entity in sorted(result.entities, key=lambda x: x.start):
            entity_text = result.original_text[entity.start:entity.end]
            writer.writerow([
                entity.entity_type,
                entity_text,
                entity.start,
                entity.end,
                f"{entity.score * 100:.1f}"
            ])
        
        return output.getvalue().encode('utf-8-sig')  # BOM for Excel
    
    # ============ DOCX EXPORTERS ============
    
    @staticmethod
    def _export_docx(result: AnalysisResult, include_metadata: bool) -> bytes:
        """Експорт анонімізованого тексту в DOCX."""
        doc = Document()
        
        # Налаштування стилів
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        if include_metadata:
            # Заголовок
            title = doc.add_heading('Анонімізований документ', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Метадані
            metadata_para = doc.add_paragraph()
            metadata_text = FileExporter._generate_metadata_header(result)
            metadata_para.add_run(metadata_text).font.size = Pt(9)
            metadata_para.add_run("\n" + "=" * 60 + "\n\n").font.size = Pt(9)
        
        # Анонімізований текст
        doc.add_paragraph(result.anonymized_text)
        
        # Зберігаємо в bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    @staticmethod
    def _export_full_report_docx(result: AnalysisResult) -> bytes:
        """Повний звіт у DOCX з форматуванням."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Звіт про деідентифікацію', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata Section
        doc.add_heading('Метадані аналізу', level=2)
        metadata_para = doc.add_paragraph(
            FileExporter._generate_metadata_header(result)
        )
        metadata_para.style.font.size = Pt(10)
        
        # Statistics
        doc.add_heading('Статистика', level=2)
        stats = FileExporter._calculate_statistics(result)
        stats_table = doc.add_table(rows=len(stats) + 1, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        # Header row
        stats_table.rows[0].cells[0].text = "Показник"
        stats_table.rows[0].cells[1].text = "Значення"
        
        # Data rows
        for idx, (key, value) in enumerate(stats.items(), 1):
            stats_table.rows[idx].cells[0].text = key
            stats_table.rows[idx].cells[1].text = str(value)
        
        doc.add_paragraph()  # Spacer
        
        # Anonymized Text
        doc.add_heading('Анонімізований текст', level=2)
        doc.add_paragraph(result.anonymized_text)
        
        doc.add_page_break()
        
        # Entities Report
        doc.add_heading('Виявлені сутності', level=2)
        
        if result.entities_count == 0:
            doc.add_paragraph("✓ Персональних даних не виявлено")
        else:
            # Групуємо за типами
            entities_by_type = {}
            for entity in result.entities:
                if entity.entity_type not in entities_by_type:
                    entities_by_type[entity.entity_type] = []
                entities_by_type[entity.entity_type].append(entity)
            
            for entity_type, entities in sorted(entities_by_type.items()):
                doc.add_heading(f'{entity_type} ({len(entities)})', level=3)
                
                for idx, entity in enumerate(sorted(entities, key=lambda x: x.start), 1):
                    entity_text = result.original_text[entity.start:entity.end]
                    para = doc.add_paragraph(style='List Number')
                    para.add_run(f"'{entity_text}' ").bold = True
                    para.add_run(
                        f"[позиція {entity.start}-{entity.end}, "
                        f"впевненість {entity.score:.0%}]"
                    ).font.size = Pt(9)
        
        # Save to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    # ============ MARKDOWN EXPORTERS ============
    
    @staticmethod
    def _export_markdown(result: AnalysisResult, include_metadata: bool) -> bytes:
        """Експорт анонімізованого тексту в Markdown."""
        lines = []
        
        if include_metadata:
            lines.extend([
                "# Анонімізований документ",
                "",
                "```",
                FileExporter._generate_metadata_header(result),
                "```",
                "",
                "---",
                ""
            ])
        
        lines.append(result.anonymized_text)
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    @staticmethod
    def _export_full_report_md(result: AnalysisResult) -> bytes:
        """Повний звіт у Markdown."""
        lines = [
            "# Звіт про деідентифікацію",
            "",
            "## Метадані аналізу",
            "",
            "```",
            FileExporter._generate_metadata_header(result),
            "```",
            "",
            "## Статистика",
            ""
        ]
        
        # Statistics table
        stats = FileExporter._calculate_statistics(result)
        lines.append("| Показник | Значення |")
        lines.append("|----------|----------|")
        for key, value in stats.items():
            lines.append(f"| {key} | {value} |")
        
        lines.extend([
            "",
            "---",
            "",
            "## Анонімізований текст",
            "",
            result.anonymized_text,
            "",
            "---",
            "",
            "## Виявлені сутності",
            ""
        ])
        
        if result.entities_count == 0:
            lines.append("✓ Персональних даних не виявлено")
        else:
            # Групуємо за типами
            entities_by_type = {}
            for entity in result.entities:
                if entity.entity_type not in entities_by_type:
                    entities_by_type[entity.entity_type] = []
                entities_by_type[entity.entity_type].append(entity)
            
            for entity_type, entities in sorted(entities_by_type.items()):
                lines.append(f"### {entity_type} ({len(entities)} знайдено)")
                lines.append("")
                
                for idx, entity in enumerate(sorted(entities, key=lambda x: x.start), 1):
                    entity_text = result.original_text[entity.start:entity.end]
                    lines.append(
                        f"{idx}. **'{entity_text}'** "
                        f"[позиція {entity.start}-{entity.end}, "
                        f"впевненість {entity.score:.0%}]"
                    )
                
                lines.append("")
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    # ============ HELPER METHODS ============
    
    @staticmethod
    def _generate_metadata_header(result: AnalysisResult) -> str:
        """Генерує header з метаданими."""
        return (
            f"Дата обробки: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Знайдено сутностей: {result.entities_count}\n"
            f"Довжина оригінального тексту: {len(result.original_text)} символів\n"
            f"Довжина анонімізованого тексту: {len(result.anonymized_text)} символів"
        )
    
    @staticmethod
    def _calculate_statistics(result: AnalysisResult) -> dict:
        """Обчислює статистику для звіту."""
        # Групуємо за типами
        entities_by_type = {}
        for entity in result.entities:
            if entity.entity_type not in entities_by_type:
                entities_by_type[entity.entity_type] = []
            entities_by_type[entity.entity_type].append(entity)
        
        # Середня впевненість
        avg_confidence = (
            sum(e.score for e in result.entities) / len(result.entities)
            if result.entities else 0
        )
        
        stats = {
            "Загальна кількість сутностей": result.entities_count,
            "Унікальних типів сутностей": len(entities_by_type),
            "Середня впевненість": f"{avg_confidence:.1%}",
            "Довжина оригінального тексту": f"{len(result.original_text)} символів",
            "Довжина анонімізованого тексту": f"{len(result.anonymized_text)} символів"
        }
        
        # Додаємо розподіл по типах
        for entity_type, entities in sorted(entities_by_type.items()):
            stats[f"  - {entity_type}"] = len(entities)
        
        return stats


# ============ CONVENIENCE FUNCTIONS ============

def generate_filename(
    base_name: str = "deidentified",
    format: str = ExportFormat.TXT,
    include_timestamp: bool = True
) -> str:
    """
    Генерує ім'я файлу для експорту.
    
    Args:
        base_name: Базова назва файлу
        format: Формат експорту
        include_timestamp: Чи додавати timestamp
        
    Returns:
        Згенероване ім'я файлу
    """
    if include_timestamp:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"{base_name}_{timestamp}.{format}"
    else:
        return f"{base_name}.{format}"