"""
Comprehensive tests для файлового I/O функціоналу.

Test Strategy:
- Unit tests для handlers та exporters
- Integration tests для повного pipeline
- Edge case coverage (encoding, corrupted files)
- Performance tests для великих файлів

Запуск: pytest test/test_file_io.py -v
"""

import pytest
import time
import tempfile
import os  # ✅ ADD THIS
from pathlib import Path
from io import BytesIO
import logging  # ✅ ADD THIS для logger в performance tests

# Initialize logger для performance tests
logger = logging.getLogger(__name__)

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.file_handlers import FileHandler, FileReadResult, sanitize_text
from utils.file_exporters import FileExporter, ExportFormat, generate_filename
from core.analyzer import AnalysisResult
from presidio_analyzer import RecognizerResult


# ============ FIXTURES ============

@pytest.fixture
def sample_txt_content():
    """Sample Ukrainian text."""
    return "Іван Петренко працює в ТОВ 'Приват'.\nEmail: ivan@example.com"


@pytest.fixture
def sample_txt_file(tmp_path, sample_txt_content):
    """Creates temporary TXT file."""
    txt_file = tmp_path / "test.txt"
    txt_file.write_text(sample_txt_content, encoding='utf-8')
    return txt_file


@pytest.fixture
def sample_docx_file(tmp_path):
    """Creates temporary DOCX file."""
    from docx import Document
    
    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Іван Петренко")
    doc.add_paragraph("Email: ivan@example.com")
    doc.save(docx_file)
    
    return docx_file


@pytest.fixture
def sample_analysis_result():
    """Creates sample AnalysisResult for export tests."""
    entities = [
        RecognizerResult(
            entity_type="PERS",
            start=0,
            end=14,
            score=0.95
        ),
        RecognizerResult(
            entity_type="EMAIL_ADDRESS",
            start=28,
            end=46,
            score=1.0
        )
    ]
    
    original_text = "Іван Петренко працює. Email: ivan@example.com"
    anonymized_text = "[PERS] працює. Email: [EMAIL_ADDRESS]"
    
    return AnalysisResult(
        entities=entities,
        anonymized_text=anonymized_text,
        original_text=original_text,
        entities_count=len(entities)
    )


# ============ FILE HANDLERS TESTS ============

class TestFileHandler:
    """Tests для FileHandler класу."""
    
    def test_read_txt_utf8(self, sample_txt_file):
        """Test: читання UTF-8 TXT файлу."""
        result = FileHandler.read_file(str(sample_txt_file))
        
        assert isinstance(result, FileReadResult)
        assert result.file_type == 'txt'
        assert result.encoding == 'utf-8'
        assert "Іван Петренко" in result.text
        assert result.char_count > 0
    
    def test_read_txt_cp1251(self, tmp_path):
        """Test: читання CP1251 файлу (fallback)."""
        cp1251_file = tmp_path / "test_cp1251.txt"
        content = "Привіт світ"
        cp1251_file.write_bytes(content.encode('cp1251'))
        
        result = FileHandler.read_file(str(cp1251_file))
        
        assert "Привіт" in result.text
        # Encoding може бути auto-detected або cp1251
        assert result.encoding in ['cp1251', 'windows-1251']
    
    def test_read_docx(self, sample_docx_file):
        """Test: читання DOCX файлу."""
        result = FileHandler.read_file(str(sample_docx_file))
        
        assert result.file_type == 'docx'
        assert "Іван Петренко" in result.text
        assert "ivan@example.com" in result.text
    
    def test_unsupported_format_raises_error(self, tmp_path):
        """Test: непідтримуваний формат викликає помилку."""
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"fake pdf content")
        
        with pytest.raises(ValueError, match="Непідтримуваний формат"):
            FileHandler.read_file(str(pdf_file))
    
    def test_file_too_large_raises_error(self, tmp_path):
        """Test: занадто великий файл."""
        large_file = tmp_path / "large.txt"
        # Симулюємо великий файл
        FileHandler.MAX_FILE_SIZE_BYTES = 100  # Temporarily reduce limit
        
        large_file.write_text("A" * 200)
        
        with pytest.raises(ValueError, match="завеликий"):
            FileHandler.read_file(str(large_file))
        
        # Restore original limit
        FileHandler.MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024
    
    def test_get_file_info(self, sample_txt_file):
        """Test: отримання метаданих файлу."""
        info = FileHandler.get_file_info(str(sample_txt_file))
        
        assert info['filename'] == 'test.txt'
        assert info['extension'] == '.txt'
        assert info['supported'] is True
        assert 'size_kb' in info


class TestSanitizeText:
    """Tests для sanitize_text функції."""
    
    def test_normalize_line_endings(self):
        """Test: нормалізація різних line endings."""
        text_windows = "Line1\r\nLine2\r\nLine3"
        text_mac = "Line1\rLine2\rLine3"
        
        result_windows = sanitize_text(text_windows)
        result_mac = sanitize_text(text_mac)
        
        assert "\r" not in result_windows
        assert "\r" not in result_mac
        assert result_windows.count("\n") == 2
    
    def test_remove_excessive_empty_lines(self):
        """Test: видалення надмірних порожніх рядків."""
        text = "Line1\n\n\n\n\nLine2"
        result = sanitize_text(text)
        
        # Повинно залишитись максимум 2 порожні рядки
        assert "\n\n\n" not in result
    
    def test_strip_trailing_whitespace(self):
        """Test: видалення пробілів в кінці рядків."""
        text = "Line1   \nLine2  \n"
        result = sanitize_text(text)
        
        assert result == "Line1\nLine2"


# ============ FILE EXPORTERS TESTS ============

class TestFileExporter:
    """Tests для FileExporter класу."""
    
    def test_export_anonymized_text_txt(self, sample_analysis_result):
        """Test: експорт анонімізованого тексту в TXT."""
        result_bytes = FileExporter.export_anonymized_text(
            sample_analysis_result,
            format=ExportFormat.TXT,
            include_metadata=True
        )
        
        content = result_bytes.decode('utf-8')
        
        assert "[PERS]" in content
        assert "[EMAIL_ADDRESS]" in content
        assert "Дата обробки:" in content
    
    def test_export_anonymized_text_docx(self, sample_analysis_result):
        """Test: експорт в DOCX."""
        result_bytes = FileExporter.export_anonymized_text(
            sample_analysis_result,
            format=ExportFormat.DOCX
        )
        
        # Перевіряємо що це валідний DOCX (ZIP archive)
        assert result_bytes[:2] == b'PK'  # ZIP signature
    
    def test_export_entities_json(self, sample_analysis_result):
        """Test: експорт сутностей в JSON."""
        result_bytes = FileExporter.export_entities_report(
            sample_analysis_result,
            format=ExportFormat.JSON
        )
        
        import json
        data = json.loads(result_bytes.decode('utf-8'))
        
        assert data['metadata']['total_entities'] == 2
        assert len(data['entities']) == 2
        assert 'statistics' in data
    
    def test_export_entities_csv(self, sample_analysis_result):
        """Test: експорт сутностей в CSV."""
        result_bytes = FileExporter.export_entities_report(
            sample_analysis_result,
            format=ExportFormat.CSV
        )
        
        content = result_bytes.decode('utf-8-sig')
        
        assert "Тип сутності" in content
        assert "PERS" in content
        assert "EMAIL_ADDRESS" in content
    
    def test_export_full_report_docx(self, sample_analysis_result):
        """Test: повний звіт в DOCX."""
        result_bytes = FileExporter.export_full_report(
            sample_analysis_result,
            format=ExportFormat.DOCX
        )
        
        # Валідний DOCX
        assert result_bytes[:2] == b'PK'
        
        # Можна додатково парсити та перевіряти вміст
        from docx import Document
        from io import BytesIO
        
        doc = Document(BytesIO(result_bytes))
        text = "\n".join(p.text for p in doc.paragraphs)
        
        assert "Звіт про деідентифікацію" in text
        assert "Метадані" in text
    
    def test_export_markdown(self, sample_analysis_result):
        """Test: експорт в Markdown."""
        result_bytes = FileExporter.export_anonymized_text(
            sample_analysis_result,
            format=ExportFormat.MARKDOWN,
            include_metadata=True
        )
        
        content = result_bytes.decode('utf-8')
        
        assert "# Анонімізований документ" in content
        assert "```" in content  # Code blocks для metadata


class TestGenerateFilename:
    """Tests для generate_filename функції."""
    
    def test_basic_filename(self):
        """Test: генерація базової назви."""
        filename = generate_filename(
            base_name="test",
            format=ExportFormat.TXT,
            include_timestamp=False
        )
        
        assert filename == "test.txt"
    
    def test_filename_with_timestamp(self):
        """Test: назва з timestamp."""
        filename = generate_filename(
            base_name="report",
            format=ExportFormat.DOCX,
            include_timestamp=True
        )
        
        assert filename.startswith("report_")
        assert filename.endswith(".docx")
        assert len(filename) > len("report_.docx")


# ============ INTEGRATION TESTS ============

@pytest.mark.integration
class TestFileIOIntegration:
    """Інтеграційні тести повного pipeline."""
    
    def test_full_pipeline_txt(self, tmp_path, sample_txt_content):
        """Test: повний цикл TXT → аналіз → експорт."""
        # 1. Створюємо файл
        input_file = tmp_path / "input.txt"
        input_file.write_text(sample_txt_content, encoding='utf-8')
        
        # 2. Читаємо файл
        read_result = FileHandler.read_file(str(input_file))
        assert read_result.text == sample_txt_content
        
        # 3. Симулюємо аналіз (тут би викликали HybridAnalyzer)
        # Для тесту використовуємо mock result
        analysis_result = AnalysisResult(
            entities=[],
            anonymized_text=read_result.text,
            original_text=read_result.text,
            entities_count=0
        )
        
        # 4. Експортуємо результат
        export_bytes = FileExporter.export_anonymized_text(
            analysis_result,
            format=ExportFormat.TXT
        )
        
        # 5. Зберігаємо експортований файл
        output_file = tmp_path / "output.txt"
        output_file.write_bytes(export_bytes)
        
        # 6. Перевіряємо
        assert output_file.exists()
        content = output_file.read_text(encoding='utf-8')
        assert sample_txt_content in content


# ============ PERFORMANCE TESTS ============


# test/test_file_io.py
# ВИДАЛІТИ benchmark параметри з performance tests

@pytest.mark.slow
class TestPerformance:
    """
    Performance tests для великих файлів.
    
    Design Philosophy: Simple timing assertions замість benchmark fixtures.
    Sufficient для smoke testing performance, extensible для production monitoring.
    """
    
    def test_large_txt_file_performance(self, tmp_path):
        """
        Test: читання великого TXT файлу виконується за прийнятний час.
        
        Performance Baseline: 10MB файл має читатись < 2 секунди.
        """
        # Arrange: Створюємо 10MB файл
        large_file = tmp_path / "large.txt"
        content = "Тестовий текст з українськими символами.\n" * 100_000
        large_file.write_text(content, encoding='utf-8')
        
        file_size_mb = large_file.stat().st_size / (1024 * 1024)
        logger.info(f"Test file size: {file_size_mb:.2f} MB")
        
        # Act: Вимірюємо час читання
        start_time = time.time()
        result = FileHandler.read_file(str(large_file))
        elapsed_time = time.time() - start_time
        
        # Assert: Performance baseline
        assert len(result.text) > 1_000_000, "File should be substantial"
        assert elapsed_time < 2.0, (
            f"Large file read took {elapsed_time:.2f}s, "
            f"expected < 2.0s for {file_size_mb:.2f}MB"
        )
        
        logger.info(
            f"Performance: Read {file_size_mb:.2f}MB in {elapsed_time:.3f}s "
            f"({file_size_mb/elapsed_time:.1f} MB/s)"
        )
    
    def test_docx_export_performance(self, sample_analysis_result):
        """
        Test: швидкість експорту в DOCX.
        
        Performance Baseline: Full report export < 3 секунди.
        """
        # Act: Вимірюємо час експорту
        start_time = time.time()
        result_bytes = FileExporter.export_full_report(
            sample_analysis_result,
            format=ExportFormat.DOCX
        )
        elapsed_time = time.time() - start_time
        
        # Assert: Performance baseline + validity
        assert len(result_bytes) > 1000, "DOCX should be substantial"
        assert result_bytes[:2] == b'PK', "Should be valid ZIP/DOCX"
        assert elapsed_time < 3.0, (
            f"DOCX export took {elapsed_time:.2f}s, expected < 3.0s"
        )
        
        logger.info(
            f"Performance: Exported DOCX in {elapsed_time:.3f}s "
            f"({len(result_bytes)/1024:.1f} KB)"
        )
    
    @pytest.mark.skipif(
        os.getenv("SKIP_STRESS_TESTS") == "1",
        reason="Stress test disabled (set SKIP_STRESS_TESTS=0 to enable)"
    )
    def test_memory_usage_large_file(self, tmp_path):
        """
        OPTIONAL: Memory stress test (skip by default).
        
        Validates memory doesn't explode with large files.
        Run explicitly: pytest -v -k memory
        """
        import psutil
        import os as os_module  # ✅ Renamed to avoid conflict
        
        # Arrange: 50MB файл
        huge_file = tmp_path / "huge.txt"
        content = "X" * (50 * 1024 * 1024)
        huge_file.write_text(content, encoding='utf-8')
        
        # Measure baseline memory
        process = psutil.Process(os_module.getpid())
        baseline_memory = process.memory_info().rss / (1024 * 1024)  # MB
        
        # Act: Read file
        result = FileHandler.read_file(str(huge_file))
        
        # Measure peak memory
        peak_memory = process.memory_info().rss / (1024 * 1024)
        memory_increase = peak_memory - baseline_memory
        
        # Assert: Memory increase reasonable (< 2x file size)
        assert memory_increase < 100, (
            f"Memory increased by {memory_increase:.1f}MB for 50MB file, "
            f"expected < 100MB"
        )
        
        logger.info(
            f"Memory: Baseline {baseline_memory:.1f}MB, "
            f"Peak {peak_memory:.1f}MB (+{memory_increase:.1f}MB)"
        )

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])